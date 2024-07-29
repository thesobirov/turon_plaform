from docx.api import Document
from docx.shared import Pt
from docx import Document
import paragraphs

document = Document('test.docx')

# qanaqadur textni topiw ucun

# for p in document.paragraphs:
#     if p.style.name.startswith("Heading") or p.style.name == "Title":
#         print(p.text)


# tableni icidan ma'lumot ovolish
# for table in document.tables:
#     print("New table")
#     for row in table.rows:
#         print("|".join([cell.text for cell in row.cells]))

#  worddagi xama textlani ovolish
# all_text = ""
# for p in document.paragraphs:
#     all_text += p.text
#     all_text += '\n'
#
# print(all_text)

#
# all_16pt_text = ""
#
# for p in document.paragraphs:
#     for run in p.runs:
#         if run.font.size == Pt(16):
#             all_16pt_text += p.text
#             all_16pt_text += "\n"
#
# print(all_16pt_text)
#
# # word fayl yaratib ishlatib korish
# document = Document()
# # h1 qowi berad
# document.add_heading("Hello world", 0)
# # p qowi berad
# p = document.add_paragraph('This is a simple text!')
#
# # shrift ozgartirish
# p.add_run(' This is a bold').bold = True
# p.add_run(' This is a italic').italic = True
# # list ul yaratad
# document.add_paragraph('This is item one', style='List Bullet')
# document.add_paragraph('This is item two', style='List Bullet')
# document.add_paragraph('This is item three', style='List Bullet')
# document.add_paragraph('This is item four', style='List Bullet')
# document.add_paragraph('This is item five', style='List Bullet')
#
# table_header = ["Name", "Age", "Job"]
#
# some_data = [
#     ["John", 46, "Programmer"],
#     ["Mary", 55, "Programmer"],
#     ["Anna", 27, "Accountant"],
#     ["Bob", 50, "Chef"],
# ]
#
# table = document.add_table(rows=1, cols=3)
# for i in range(3):
#     table.rows[0].cells[i].text = table_header[i]
#
# for name, age, job in some_data:
#     cells = table.add_row().cells
#     cells[0].text = name
#     cells[1].text = str(age)
#     cells[2].text = job
#
# document.add_page_break()
#
# document.add_paragraph("HELLO NEW PAGE")
#
# document.add_picture('neuronal.jpg')
#
# document.save('test.docx')
